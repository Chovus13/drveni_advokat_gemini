# extract_and_structure.py

import os
import re
import json
import argparse
import docx
from tqdm import tqdm
import logging
import config  # << NOVO: Uvozimo naš konfiguracioni fajl


# Podešavanje za logovanje grešaka
logging.basicConfig(
    filename='extraction_log.txt',
    level=logging.WARNING,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def extract_and_clean_document(document: docx.Document) -> tuple[str, dict]:
    """
    NOVA I POBOLJŠANA FUNKCIJA
    Prima ceo docx dokument objekat i iz njega izdvaja:
    1. Očišćen tekst (bez zaglavlja, podnožja i boilerplate fraza).
    2. Ekstrahovane metapodatke.
    """
    
    # --- DEO ZA ČIŠĆENJE TEKSTA ---
    
    # Prvo, prikupljamo sav tekst koji treba ukloniti
    text_to_remove = set() # Koristimo 'set' da izbegnemo duplikate
    
    # 1. Uklanjanje na osnovu konfiguracione liste
    for phrase in config.BOILERPLATE_PHRASES_TO_REMOVE:
        text_to_remove.add(phrase)

    # 2. Direktno čitanje i uklanjanje teksta iz header-a i footer-a
    if config.REMOVE_HEADERS_FOOTERS:
        for section in document.sections:
            # Zaglavlja
            for para in section.header.paragraphs:
                if para.text:
                    text_to_remove.add(para.text.strip())
            # Podnožja
            for para in section.footer.paragraphs:
                if para.text:
                    text_to_remove.add(para.text.strip())

    # Sastavljanje glavnog teksta iz paragrafa
    main_text = "\n".join([para.text for para in document.paragraphs])
    
    # Uklanjanje prikupljenih fraza iz glavnog teksta
    for phrase_to_remove in text_to_remove:
        if phrase_to_remove: # Ne uklanjamo prazne stringove
            main_text = main_text.replace(phrase_to_remove, "")
            
    # Uklanjanje višestrukih praznih redova nakon čišćenja
    main_text = re.sub(r'\n{2,}', '\n', main_text).strip()
    
    # --- DEO ZA EKSTRAKCIJU METAPODATAKA (ostaje skoro isti) ---
    metadata = extract_metadata_from_text(main_text)
    
    return main_text, metadata


def extract_metadata_from_text(text: str) -> dict:
    """
    Koristi regularne izraze za ekstrakciju metapodataka iz teksta dokumenta.
    Dizajnirano da bude fleksibilno za varijacije u formatiranju.
    """
    metadata = {}
    
    # Definicija regularnih izraza za svaki metapodatak.
    # re.IGNORECASE čini pretragu neosetljivom na velika/mala slova.
    # re.DOTALL omogućava da '.' obuhvati i nove redove, korisno za višeredne unose.
    patterns = {
        "case_id": r"Broj predmeta:?\s*([\w\d\s\/-]+)",
        "judge": r"Sudija:?\s*([A-ZŠĐČĆŽ][a-zšđčćž]+(?:\s+[A-ZŠĐČĆŽ][a-zšđčćž]+)+)",
        "plaintiff": r"Tužilac:?\s*(.*?)(?=\nTuženi:|Sudija:)",
        "defendant": r"Tuženi:?\s*(.*?)(?=\nSud:|Datum presude:)",
        "court": r"Sud:?\s*(.*?)(?=\n|$)",
        "decision_date": r"Datum presude:?\s*(\d{1,2}\.\d{1,2}\.\d{4}\.?|\d{4}-\d{2}-\d{2})",
        # Primer za tip dokumenta - traži reč "PRESUDA" ili "REŠENJE" velikim slovima
        "document_type": r"\b(PRESUDA|REŠENJE)\b" 
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            # .strip() uklanja praznine sa početka i kraja
            extracted_value = match.group(1).strip()
            
            # Specijalna obrada za tužene, koji mogu biti lista
            if key in ["defendant", "plaintiff"]:
                # Deli na osnovu novog reda ili zareza i čisti svaki unos
                items = re.split(r'[\n,]+', extracted_value)
                metadata[key] = [item.strip() for item in items if item.strip()]
            else:
                metadata[key] = extracted_value

    return metadata

def clean_full_text(text: str) -> str:
    """
    Osnovno čišćenje teksta. Uklanja višestruke prazne redove.
    Ovde se mogu dodati složenija pravila za uklanjanje zaglavlja/podnožja.
    """
    # Zameni dva ili više novih redova sa jednim
    text = re.sub(r'\n{2,}', '\n', text)
    return text.strip()

def process_docx_files(source_dir: str, output_path: str):
    """
    Glavna funkcija koja obrađuje sve .docx fajlove u direktorijumu.
    
    Glavna funkcija - sada poziva novu, pametniju funkciju za obradu.
    """
    print(f"Započinjanje ekstrakcije iz direktorijuma: {source_dir}")
    
    # Prikupljanje svih .docx fajlova za prikaz status bara
    docx_files = [
        os.path.join(root, file)
        for root, _, files in os.walk(source_dir)
        for file in files
        if file.lower().endswith(".docx")
    ]

    if not docx_files:
        print("Nema .docx fajlova u navedenom direktorijumu.")
        return

    # Otvaranje izlaznog .jsonl fajla sa UTF-8 kodiranjem
    with open(output_path, 'w', encoding='utf-8') as outfile:
        # Malo ispravljena logika za pronalaženje fajlova unutar tqdm
        all_files = [os.path.join(root, file) for root, _, files in os.walk(source_dir) for file in files if file.lower().endswith('.docx')]
        for file_path in tqdm(all_files, desc="Procesiranje dokumenata"):
            try:
                document = docx.Document(file_path)
                
                # << JEDINA KLJUČNA IZMENA U OVOJ FUNKCIJI >>
                # Pozivamo novu funkciju koja radi i čišćenje i ekstrakciju
                cleaned_text, metadata = extract_and_clean_document(document)
                
                # Kreiranje finalnog JSON objekta (logika ostaje ista)
                structured_data = {
                    "source_file": file_path,
                    "case_id": metadata.get("case_id", "Nepoznato"),
                    "full_text": cleaned_text,
                    "metadata": {
                        "judge": metadata.get("judge", "Nepoznato"),
                        "plaintiff": metadata.get("plaintiff", []),
                        "defendant": metadata.get("defendant", []),
                        "decision_date": metadata.get("decision_date", "Nepoznato"),
                        "court": metadata.get("court", "Nepoznato"),
                        "document_type": metadata.get("document_type", "Nepoznato")
                    }
                }
                
                json.dump(structured_data, outfile, ensure_ascii=False)
                outfile.write('\n')

            except Exception as e:
                error_message = f"Greška pri obradi fajla {file_path}: {e}"
                # print(error_message) # Možemo isključiti ispis u konzoli da ne bude pretrpano
                logging.warning(error_message)

    print(f"\nEkstrakcija završena. Podaci sačuvani u: {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Ekstrahuje tekst i metapodatke iz .docx fajlova i čuva ih u JSONL formatu."
    )
    parser.add_argument(
        "source_directory",
        type=str,
        help="Putanja do direktorijuma sa konvertovanim .docx fajlovima."
    )
    parser.add_argument(
        "output_file",
        type=str,
        help="Putanja do izlaznog .jsonl fajla."
    )

    args = parser.parse_args()

    # Provera da li postoji `tqdm` i davanje instrukcija ako ne postoji
    try:
        from tqdm import tqdm
    except ImportError:
        print("Biblioteka 'tqdm' nije instalirana. Preporučuje se instalacija za prikaz statusa.")
        print("Instalirajte je komandom: pip install tqdm")
        # Definišemo 'dummy' tqdm funkciju ako ne postoji, da skripta ne pukne
        tqdm = lambda x, **kwargs: x

    process_docx_files(args.source_directory, args.output_file)